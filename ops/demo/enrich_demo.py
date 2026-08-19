#!/usr/bin/env python3
from __future__ import annotations

import random
from datetime import date, timedelta
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app import create_app
from app.extensions import db
from app.models import (
    ClinicSettings,
    DocumentDefinition,
    DocumentField,
    Employee,
    EmployeeAbsence,
    PatientAccountingCategory,
    PatientAccountingHistory,
    PatientAccountingPatient,
    PatientDispensaryObservation,
    PatientPassportMembership,
    Position,
    User,
    VacationPlan,
)
from app.modules.patient_accounting.service import get_origin_for_user


DOCUMENT_KEY = "demo_vacation_statement"
TEMPLATE_FILENAME = "demo_vacation_statement.docx"


def _employee_number(employee: Employee) -> int | None:
    value = str(employee.personnel_number or "")
    if not value.startswith("DEMO-"):
        return None

    try:
        return int(value.split("-", 1)[1])
    except (IndexError, ValueError):
        return None


def _scattered_vacation_starts(
    *,
    year: int,
    employee_number: int,
    today: date,
    force_current: bool,
) -> list[date]:
    """Three deterministic, staggered 14-day vacation starts."""
    year_start = date(year, 1, 1)
    latest_start = date(year, 12, 18)
    latest_offset = (latest_start - year_start).days

    anchors = (38, 158, 278)
    rng = random.Random(
        year * 100_003
        + employee_number * 7_919
        + 106
    )

    offsets = [
        max(
            0,
            min(
                latest_offset,
                anchor + rng.randint(-48, 48),
            ),
        )
        for anchor in anchors
    ]

    if force_current and today.year == year:
        current_start = today - timedelta(days=3)
        if current_start < year_start:
            current_start = year_start
        if current_start > latest_start:
            current_start = latest_start

        current_offset = (current_start - year_start).days

        replace_index = min(
            range(3),
            key=lambda index: abs(
                offsets[index] - current_offset
            ),
        )
        offsets[replace_index] = current_offset

    starts = sorted(
        year_start + timedelta(days=offset)
        for offset in offsets
    )

    normalized: list[date] = []
    for start in starts:
        if normalized:
            minimum = normalized[-1] + timedelta(days=21)
            if start < minimum:
                start = minimum

        if start > latest_start:
            start = latest_start

        normalized.append(start)

    for index in range(len(normalized) - 2, -1, -1):
        maximum = normalized[index + 1] - timedelta(days=21)
        if normalized[index] > maximum:
            normalized[index] = maximum

    return normalized


def _randomize_vacations(today: date) -> int:
    plans = (
        VacationPlan.query
        .filter_by(
            year=today.year,
            status="approved",
        )
        .all()
    )

    changed_periods = 0

    for plan in plans:
        employee = plan.employee
        if employee is None:
            continue

        number = _employee_number(employee)
        if number is None:
            continue

        force_current = number in {4, 18}

        starts = _scattered_vacation_starts(
            year=today.year,
            employee_number=number,
            today=today,
            force_current=force_current,
        )

        periods = sorted(
            plan.approved_periods,
            key=lambda period: (
                period.sort_order,
                period.id or 0,
            ),
        )

        for sort_order, (period, start) in enumerate(
            zip(periods[:3], starts),
            start=1,
        ):
            period.start_date = start
            period.days_count = 14
            period.end_date = start + timedelta(days=13)
            period.sort_order = sort_order
            changed_periods += 1

        if force_current:
            current_period = next(
                (
                    period
                    for period in periods
                    if period.start_date <= today <= period.end_date
                ),
                None,
            )
            if current_period is not None:
                absences = (
                    EmployeeAbsence.query
                    .filter_by(employee_id=employee.id)
                    .all()
                )
                for absence in absences:
                    if (
                        absence.absence_type
                        and absence.absence_type.name == "Отпуск"
                    ):
                        absence.start_date = current_period.start_date
                        absence.end_date = current_period.end_date

    return changed_periods


def _apply_demo_cases() -> None:
    employee = (
        Employee.query
        .filter_by(personnel_number="DEMO-001")
        .first()
    )

    if employee is not None:
        employee.full_name_genitive = "Иванова Ивана Ивановича"
        employee.full_name_dative = "Иванову Ивану Ивановичу"

    position = (
        Position.query
        .filter_by(name="Заведующий терапевтическим отделением")
        .first()
    )

    if position is not None:
        position.name_genitive = (
            "заведующего терапевтическим отделением"
        )
        position.name_dative = (
            "заведующему терапевтическим отделением"
        )


def _write_vacation_template(*, target_path: Path) -> None:
    document = DocxDocument()

    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    recipient_lines = (
        "Главному врачу",
        "Демонстрационной городской поликлиники PZ-Med",
        "{{ clinic_head_full_name_dative }}",
        "от {{ author_position_genitive }}",
        "{{ author_full_name_genitive }}",
    )

    for line in recipient_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(9)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(line)

    document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run("ЗАЯВЛЕНИЕ")
    run.bold = True
    run.font.size = Pt(14)

    body = document.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.first_line_indent = Cm(1.25)
    body.paragraph_format.line_spacing = 1.15
    body.add_run(
        "Прошу предоставить мне ежегодный оплачиваемый отпуск с "
        "{{ vacation_start_date }} продолжительностью "
        "{{ vacation_days_calendar_days_text }}."
    )

    document.add_paragraph()

    signature = document.add_table(rows=1, cols=3)
    signature.autofit = False
    signature.columns[0].width = Cm(4)
    signature.columns[1].width = Cm(5)
    signature.columns[2].width = Cm(5)

    cells = signature.rows[0].cells
    cells[0].text = "{{ current_date }}"
    cells[1].text = "________________"
    cells[2].text = "{{ author_short_name }}"

    for cell in cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

    target_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(target_path)


def _ensure_vacation_document(app) -> None:
    system_user = (
        User.query
        .filter_by(username="demo_system_internal")
        .first()
    )

    document = (
        DocumentDefinition.query
        .filter_by(document_key=DOCUMENT_KEY)
        .first()
    )

    if document is None:
        document = DocumentDefinition(
            document_key=DOCUMENT_KEY,
            title="Заявление на отпуск",
            document_type="statement",
        )
        db.session.add(document)

    document.title = "Заявление на отпуск"
    document.document_type = "statement"
    document.description = (
        "Демонстрационный шаблон заявления "
        "на ежегодный оплачиваемый отпуск."
    )
    document.template_filename = TEMPLATE_FILENAME
    document.created_by_user_id = (
        system_user.id
        if system_user is not None
        else None
    )
    document.is_active = True
    document.all_subdivisions = True
    document.all_departments = True
    document.all_positions = True
    document.all_roles = True
    document.set_allowed_roles([])

    db.session.flush()

    desired_fields = (
        {
            "field_key": "vacation_start_date",
            "label": "Дата начала отпуска",
            "field_type": "date",
            "placeholder": "",
            "help_text": (
                "Укажите первый день ежегодного оплачиваемого отпуска."
            ),
            "default_source": "none",
            "is_required": True,
            "sort_order": 10,
        },
        {
            "field_key": "vacation_days",
            "label": "Количество дней",
            "field_type": "number",
            "placeholder": "14",
            "help_text": (
                "Укажите продолжительность отпуска в календарных днях."
            ),
            "default_source": "none",
            "is_required": True,
            "sort_order": 20,
        },
    )

    existing = {
        field.field_key: field
        for field in document.fields
    }

    for values in desired_fields:
        field = existing.get(values["field_key"])

        if field is None:
            field = DocumentField(
                document_id=document.id,
                field_key=values["field_key"],
                label=values["label"],
            )
            db.session.add(field)

        field.label = values["label"]
        field.field_type = values["field_type"]
        field.placeholder = values["placeholder"]
        field.help_text = values["help_text"]
        field.default_source = values["default_source"]
        field.is_required = values["is_required"]
        field.sort_order = values["sort_order"]

    template_folder = Path(
        app.config["DOCUMENT_TEMPLATE_FOLDER"]
    )
    _write_vacation_template(
        target_path=template_folder / TEMPLATE_FILENAME
    )


def _ensure_patient_category(
    *,
    kind: str,
    name: str,
    description: str,
    default_interval_days: int | None,
) -> PatientAccountingCategory:
    category = (
        PatientAccountingCategory.query
        .filter_by(kind=kind, name=name)
        .first()
    )
    if category is None:
        category = PatientAccountingCategory(
            kind=kind,
            name=name,
        )
        db.session.add(category)

    category.description = description
    category.default_interval_days = (
        default_interval_days
        if kind == "observation"
        else None
    )
    category.is_active = True
    db.session.flush()
    return category


def _ensure_patient_card(
    *,
    owner: User,
    card_number: str,
    comment: str,
) -> tuple[PatientAccountingPatient, bool]:
    key = card_number.strip().casefold()
    patient = (
        PatientAccountingPatient.query
        .filter_by(card_number_key=key)
        .first()
    )
    created = patient is None

    subdivision_id, department_id = get_origin_for_user(owner)

    if patient is None:
        patient = PatientAccountingPatient(
            card_number=card_number,
            card_number_key=key,
            responsible_user_id=owner.id,
            source_subdivision_id=subdivision_id,
            source_department_id=department_id,
            created_by_user_id=owner.id,
        )
        db.session.add(patient)
        db.session.flush()
    else:
        patient.responsible_user_id = owner.id
        patient.source_subdivision_id = subdivision_id
        patient.source_department_id = department_id

    patient.comment = comment
    patient.is_active = True
    return patient, created


def _add_demo_history(
    *,
    patient: PatientAccountingPatient,
    owner: User,
    event_type: str,
    message: str,
) -> None:
    exists = (
        PatientAccountingHistory.query
        .filter_by(
            patient_id=patient.id,
            event_type=event_type,
            message=message,
        )
        .first()
    )
    if exists is None:
        db.session.add(
            PatientAccountingHistory(
                patient_id=patient.id,
                actor_user_id=owner.id,
                event_type=event_type,
                message=message,
            )
        )


def _ensure_demo_patient_accounting(today: date) -> int:
    owner = User.query.filter_by(username="demo_ivanov").first()
    if owner is None:
        raise RuntimeError("Demo patient accounting: demo_ivanov not found")

    passport_svo = _ensure_patient_category(
        kind="passport",
        name="участники СВО",
        description=(
            "Демонстрационная категория паспорта участка для контроля "
            "пациентов, относящихся к участникам СВО."
        ),
        default_interval_days=None,
    )
    passport_mobile = _ensure_patient_category(
        kind="passport",
        name="Маломобильные граждане",
        description=(
            "Демонстрационная категория пациентов, которым может "
            "потребоваться доступная маршрутизация и помощь на дому."
        ),
        default_interval_days=None,
    )
    diabetes = _ensure_patient_category(
        kind="observation",
        name="Сахарный диабет",
        description="Плановое диспансерное наблюдение каждые 90 дней.",
        default_interval_days=90,
    )
    hypertension = _ensure_patient_category(
        kind="observation",
        name="Гипертоническая болезнь",
        description="Плановое диспансерное наблюдение каждые 180 дней.",
        default_interval_days=180,
    )

    created_cards = 0

    patient, created = _ensure_patient_card(
        owner=owner,
        card_number="118742",
        comment=(
            "Демонстрационная запись. Категория паспорта участка "
            "подтверждена, контроль сведений при плановом обращении."
        ),
    )
    created_cards += int(created)
    membership = PatientPassportMembership.query.filter_by(
        patient_id=patient.id,
        category_id=passport_svo.id,
    ).first()
    if membership is None:
        db.session.add(
            PatientPassportMembership(
                patient_id=patient.id,
                category_id=passport_svo.id,
                added_by_user_id=owner.id,
            )
        )
    _add_demo_history(
        patient=patient,
        owner=owner,
        event_type="demo_seeded",
        message="Пациент включён в категорию паспорта участка «участники СВО».",
    )

    patient, created = _ensure_patient_card(
        owner=owner,
        card_number="204615",
        comment=(
            "Демонстрационная запись. При обращении учитывать "
            "маломобильность и необходимость доступной маршрутизации."
        ),
    )
    created_cards += int(created)
    membership = PatientPassportMembership.query.filter_by(
        patient_id=patient.id,
        category_id=passport_mobile.id,
    ).first()
    if membership is None:
        db.session.add(
            PatientPassportMembership(
                patient_id=patient.id,
                category_id=passport_mobile.id,
                added_by_user_id=owner.id,
            )
        )
    _add_demo_history(
        patient=patient,
        owner=owner,
        event_type="demo_seeded",
        message=(
            "Пациент включён в категорию паспорта участка "
            "«Маломобильные граждане»."
        ),
    )

    patient, created = _ensure_patient_card(
        owner=owner,
        card_number="317506",
        comment=(
            "Демонстрационная запись. Плановый контроль показателей "
            "углеводного обмена и соблюдения рекомендаций."
        ),
    )
    created_cards += int(created)
    observation = PatientDispensaryObservation.query.filter_by(
        patient_id=patient.id,
        category_id=diabetes.id,
    ).first()
    if observation is None:
        observation = PatientDispensaryObservation(
            patient_id=patient.id,
            category_id=diabetes.id,
            added_by_user_id=owner.id,
        )
        db.session.add(observation)
    observation.comment = "Плановый контроль; требуется пригласить пациента."
    patient.observation_status = "invite"
    patient.next_invite_date = today
    patient.completed_on = None
    _add_demo_history(
        patient=patient,
        owner=owner,
        event_type="demo_seeded",
        message="Добавлено диспансерное наблюдение «Сахарный диабет».",
    )

    patient, created = _ensure_patient_card(
        owner=owner,
        card_number="428193",
        comment=(
            "Демонстрационная запись. Контроль артериального давления, "
            "терапии и факторов сердечно-сосудистого риска."
        ),
    )
    created_cards += int(created)
    observation = PatientDispensaryObservation.query.filter_by(
        patient_id=patient.id,
        category_id=hypertension.id,
    ).first()
    if observation is None:
        observation = PatientDispensaryObservation(
            patient_id=patient.id,
            category_id=hypertension.id,
            added_by_user_id=owner.id,
        )
        db.session.add(observation)
    observation.comment = "Очередной цикл наблюдения пройден без замечаний."
    patient.observation_status = "completed"
    patient.completed_on = today - timedelta(days=45)
    patient.next_invite_date = patient.completed_on + timedelta(days=180)
    _add_demo_history(
        patient=patient,
        owner=owner,
        event_type="demo_seeded",
        message="Добавлено диспансерное наблюдение «Гипертоническая болезнь».",
    )

    db.session.flush()
    return created_cards


def _configure_demo_clinic_head() -> None:
    clinic = (
        ClinicSettings.query
        .order_by(ClinicSettings.id.asc())
        .first()
    )

    if clinic is None:
        return

    clinic.head_name = "Соколов Андрей Викторович"
    clinic.head_name_genitive = "Соколова Андрея Викторовича"
    clinic.head_name_dative = "Соколову Андрею Викторовичу"
    clinic.head_position = "главный врач"
    clinic.head_position_genitive = "главного врача"
    clinic.head_position_dative = "главному врачу"


def enrich() -> None:
    app = create_app()
    today = date.today()

    with app.app_context():
        changed_periods = _randomize_vacations(today)
        _apply_demo_cases()
        _configure_demo_clinic_head()
        _ensure_vacation_document(app)
        patient_cards = _ensure_demo_patient_accounting(today)
        db.session.commit()

        print(
            "Demo enrich complete: "
            f"{changed_periods} vacation periods randomized"
        )
        print("Demo document ready: Заявление на отпуск")
        print(
            "Demo patient accounting ready: "
            "2 passport categories, 2 observation categories, "
            f"4 cards ({patient_cards} newly created)"
        )


if __name__ == "__main__":
    enrich()
